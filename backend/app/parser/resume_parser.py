# app/parser/resume_parser.py
"""
Resume Parser Module
====================
Supports : PDF, DOCX, TXT
Extracts : name, email, phone, skills, experience, education

NLP pipeline (in priority order):
  1. spaCy en_core_web_sm  — Named Entity Recognition (PERSON, ORG, DATE)
  2. Regex patterns         — email, phone, experience years, education level
  3. Heuristic fallbacks    — first-line name detection, date-range maths
"""

import re
import os
import logging
import datetime
from pathlib import Path
from typing import Optional, List

logger = logging.getLogger(__name__)

# ── spaCy lazy loader ─────────────────────────────────────────────────────────
_nlp = None
_SPACY_OK = False

def _get_nlp():
    global _nlp, _SPACY_OK
    if _nlp is not None:
        return _nlp
    try:
        import spacy
        _nlp = spacy.load("en_core_web_sm")
        _SPACY_OK = True
        logger.info("spaCy model 'en_core_web_sm' loaded")
    except OSError:
        logger.warning("spaCy model not found. Run: python -m spacy download en_core_web_sm")
        _nlp = None; _SPACY_OK = False
    except ImportError:
        logger.warning("spaCy not installed. Using regex-only parsing.")
        _nlp = None; _SPACY_OK = False
    return _nlp

# ── Text extraction ───────────────────────────────────────────────────────────

def extract_text_from_pdf(path):
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t: text += t + "\n"
        if text.strip(): return text
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}")
    try:
        import PyPDF2
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ""
    except Exception as e:
        logger.error(f"PyPDF2 failed: {e}")
    return text

def extract_text_from_docx(path):
    try:
        from docx import Document
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip(): parts.append(cell.text)
        return "\n".join(parts)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}"); return ""

def extract_text_from_txt(path):
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f: return f.read()
    except Exception as e:
        logger.error(f"TXT failed: {e}"); return ""

def extract_text(file_path):
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":  return extract_text_from_pdf(file_path)
    elif ext == ".docx": return extract_text_from_docx(file_path)
    elif ext == ".txt":  return extract_text_from_txt(file_path)
    else: raise ValueError(f"Unsupported format '{ext}'. Use .pdf .docx .txt")

# ── spaCy NER helpers ─────────────────────────────────────────────────────────

def _spacy_doc(text):
    nlp = _get_nlp()
    return nlp(text[:5000]) if nlp else None

def extract_name_spacy(text):
    """Use spaCy PERSON entity in the first 1500 chars."""
    nlp = _get_nlp()
    if not nlp: return None
    doc = nlp(text[:1500])
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            words = ent.text.strip().split()
            if 2 <= len(words) <= 4 and all(w.isalpha() for w in words):
                return ent.text.strip().title()
    return None

def extract_orgs_spacy(text):
    """Return list of ORG entity strings (employers, universities)."""
    doc = _spacy_doc(text)
    if not doc: return []
    return list({e.text.strip() for e in doc.ents if e.label_ == "ORG"})

def extract_dates_spacy(text):
    """Return list of DATE entity strings."""
    doc = _spacy_doc(text)
    if not doc: return []
    return [e.text.strip() for e in doc.ents if e.label_ == "DATE"]

# ── Regex extractors ──────────────────────────────────────────────────────────

def extract_email(text):
    m = re.findall(r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b', text)
    return m[0] if m else ""

def extract_phone(text):
    for pat in [r'\+?1?\s*[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', r'\+?\d[\d\s\-().]{8,}\d']:
        m = re.findall(pat, text)
        if m: return m[0].strip()
    return ""

def extract_name_heuristic(text):
    labelled = re.search(
        r'(?:^|\n)\s*(?:name|full\s+name)\s*[:\-]\s*([A-Za-z][A-Za-z\s\-\.]{2,40})',
        text, re.IGNORECASE)
    if labelled: return labelled.group(1).strip().title()
    skip = {"resume","curriculum","vitae","cv","objective","summary","experience",
            "education","skills","email","phone","address","linkedin","github"}
    for line in text.split("\n")[:12]:
        line = line.strip()
        if not line or "@" in line or "http" in line: continue
        if any(kw in line.lower() for kw in skip): continue
        words = line.split()
        if 2 <= len(words) <= 4 and all(re.match(r'^[A-Za-z][\w\-\.]*$', w) for w in words):
            return line.title()
    return "Unknown Candidate"

def extract_name(text):
    """spaCy PERSON NER → heuristic fallback."""
    name = extract_name_spacy(text)
    if name: return name
    return extract_name_heuristic(text)

def extract_experience_years(text):
    """
    Four-strategy pipeline:
    1. Explicit phrases ('5 years of experience')
    2. YYYY-present date ranges
    3. spaCy DATE entity years
    4. All 4-digit years in document (min→now)
    """
    years_found = []
    cur = datetime.datetime.now().year

    # Strategy 1 — explicit phrases
    for pat in [
        r'(\d+\.?\d*)\s*\+?\s*years?\s+(?:of\s+)?(?:work\s+|professional\s+)?experience',
        r'experience\s*[:\-]?\s*(\d+\.?\d*)\s*\+?\s*years?',
        r'(\d+\.?\d*)\s*\+?\s*yrs?\s+(?:of\s+)?(?:work\s+)?experience',
        r'over\s+(\d+\.?\d*)\s*years?\s+(?:of\s+)?experience',
    ]:
        for m in re.findall(pat, text.lower()):
            try:
                v = float(m)
                if 0 < v <= 50: years_found.append(v)
            except ValueError: pass

    # Strategy 2 — YYYY-present
    for m in re.findall(r'(\b(?:19|20)\d{2}\b)\s*[-–—to]+\s*(?:present|current|now)', text.lower()):
        try:
            d = cur - int(m)
            if 0 < d <= 50: years_found.append(float(d))
        except ValueError: pass

    # Strategy 3 — spaCy DATE entities
    spacy_years = []
    for d in extract_dates_spacy(text):
        spacy_years += [int(y) for y in re.findall(r'\b(19\d{2}|20\d{2})\b', d)]
    if len(spacy_years) >= 2:
        earliest = min(spacy_years)
        if earliest < cur: years_found.append(float(cur - earliest))

    # Strategy 4 — all years in document
    if not years_found:
        all_y = [int(y) for y in re.findall(r'\b(19\d{2}|20\d{2})\b', text)]
        if len(all_y) >= 2:
            e = min(all_y)
            if e < cur: years_found.append(float(cur - e))

    return round(min(max(years_found), 50.0), 1) if years_found else 0.0

def extract_education(text):
    edu_map = [
        (["ph.d","phd","ph. d","doctorate","doctoral"], "PhD"),
        (["master of science","master of arts","master of engineering","m.s.","m.sc",
          "m.a.","m.e.","mtech","m.tech","master's","masters"], "Master's"),
        (["mba","master of business administration"], "MBA"),
        (["bachelor of science","bachelor of arts","bachelor of engineering",
          "b.s.","b.sc","b.a.","b.e.","b.tech","btech","bachelor's","bachelors"], "Bachelor's"),
        (["associate","a.s.","a.a."], "Associate's"),
        (["diploma","polytechnic"], "Diploma"),
        (["high school","secondary school","12th","hsc","ssc"], "High School"),
    ]
    tl = text.lower()
    for kws, label in edu_map:
        if any(kw in tl for kw in kws): return label
    return ""

def extract_linkedin(text):
    m = re.search(r'linkedin\.com/in/[A-Za-z0-9\-_/]+', text, re.IGNORECASE)
    return m.group(0) if m else ""

def extract_github(text):
    m = re.search(r'github\.com/[A-Za-z0-9\-_]+', text, re.IGNORECASE)
    return m.group(0) if m else ""

# ── Main entry point ──────────────────────────────────────────────────────────

def parse_resume(file_path: str) -> dict:
    """
    Full pipeline:  extract text → NER → regex → return structured dict.
    skills list is populated later by skill_extractor.py.
    """
    raw_text = extract_text(file_path)
    if not raw_text.strip():
        raise ValueError(
            "No text extracted. The file may be empty, image-based (scanned PDF), "
            "or password-protected."
        )
    result = {
        "name":          extract_name(raw_text),
        "email":         extract_email(raw_text),
        "phone":         extract_phone(raw_text),
        "experience":    extract_experience_years(raw_text),
        "education":     extract_education(raw_text),
        "linkedin":      extract_linkedin(raw_text),
        "github":        extract_github(raw_text),
        "organizations": extract_orgs_spacy(raw_text),
        "raw_text":      raw_text,
        "skills":        [],
        "spacy_used":    _SPACY_OK,
    }
    logger.info(f"Parsed '{Path(file_path).name}': {result['name']} | "
                f"exp={result['experience']}y | edu={result['education']} | "
                f"spaCy={'yes' if _SPACY_OK else 'fallback'}")
    return result
